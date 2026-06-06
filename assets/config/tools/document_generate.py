"""Document generation tool for new-file outputs."""


from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Any

from harzoo.agent.kernel.tool import Tool, ToolResult

TOOL_VERSION = "2026-05-22"


def _resolve_path(path: str) -> Path:
    p = Path(path).expanduser()
    if p.is_absolute():
        return p.resolve()
    return (Path.cwd() / p).resolve()


class DocumentGenerateTool(Tool):
    """文档生成工具：根据结构化载荷创建目标格式文档。"""

    name = "DocumentGenerate"
    danger_level = 0
    description = "Generate new documents for common formats with explicit payloads."
    parameters = {
        "type": "object",
        "properties": {
            "output_path": {"type": "string"},
            "format": {"type": "string", "enum": ["pdf", "txt", "json", "csv", "docx", "xlsx"]},
            "document": {"description": "Structured document payload"},
            "overwrite": {"type": "boolean", "default": False},
        },
        "required": ["output_path", "format", "document"],
    }
    risk_level = "medium"
    side_effect = True
    idempotent = False

    def execute(self, output_path: str, format: str, document: Any, overwrite: bool = False, **_: Any) -> ToolResult:
        fmt = str(format).strip().lower()
        if fmt not in {"pdf", "txt", "json", "csv", "docx", "xlsx"}:
            return ToolResult.failure(f"Unsupported format: {fmt}", code="UNSUPPORTED_FORMAT")
        path = _resolve_path(output_path)
        if path.exists() and not overwrite:
            return ToolResult.failure(f"Output already exists: {output_path}", code="CONFLICT")
        if not isinstance(document, dict):
            return ToolResult.failure("document must be an object", code="INVALID_ARGUMENTS")
        path.parent.mkdir(parents=True, exist_ok=True)
        try:
            if fmt == "pdf":
                blocks = document.get("blocks")
                if not isinstance(blocks, list) or not blocks:
                    return ToolResult.failure("document.blocks must be a non-empty list", code="INVALID_ARGUMENTS")
                self._generate_pdf(path, blocks=blocks)
            elif fmt == "txt":
                self._generate_txt(path, document)
            elif fmt == "json":
                self._generate_json(path, document)
            elif fmt == "csv":
                self._generate_csv(path, document)
            elif fmt == "docx":
                self._generate_docx(path, document)
            elif fmt == "xlsx":
                self._generate_xlsx(path, document)
        except Exception as exc:  # noqa: BLE001
            return ToolResult.failure(f"{type(exc).__name__}: {exc}", code="WRITE_FAILED")
        return ToolResult.success(
            {
                "path": str(path),
                "format": fmt,
                "operation": "create",
                "bytes_written": path.stat().st_size,
                "warnings": [],
            }
        )

    def _generate_txt(self, path: Path, document: dict[str, Any]) -> None:
        content = str(document.get("content", ""))
        path.write_text(content, encoding="utf-8")

    def _generate_json(self, path: Path, document: dict[str, Any]) -> None:
        payload = document.get("data")
        if payload is None:
            raise ValueError("document.data is required for json")
        path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    def _generate_csv(self, path: Path, document: dict[str, Any]) -> None:
        rows = document.get("rows")
        if not isinstance(rows, list):
            raise ValueError("document.rows must be a list")
        with path.open("w", encoding="utf-8", newline="") as fh:
            writer = csv.writer(fh)
            for row in rows:
                if not isinstance(row, list):
                    raise ValueError("each row in document.rows must be a list")
                writer.writerow([str(cell) for cell in row])

    def _generate_docx(self, path: Path, document: dict[str, Any]) -> None:
        try:
            from docx import Document  # pyright: ignore[reportMissingImports]
        except Exception as exc:  # noqa: BLE001
            raise RuntimeError(f"DOCX dependency missing: {exc}") from exc
        doc = Document()
        blocks = document.get("blocks")
        if isinstance(blocks, list):
            for block in blocks:
                if not isinstance(block, dict):
                    continue
                text = str(block.get("text") or "").strip()
                if not text:
                    continue
                if str(block.get("type") or "") == "heading":
                    doc.add_heading(text, level=1)
                else:
                    doc.add_paragraph(text)
        else:
            content = str(document.get("content", "")).strip()
            if content:
                for line in content.splitlines():
                    if line.strip():
                        doc.add_paragraph(line)
        doc.save(str(path))

    def _generate_xlsx(self, path: Path, document: dict[str, Any]) -> None:
        try:
            from openpyxl import Workbook  # pyright: ignore[reportMissingImports,reportMissingModuleSource]
        except Exception as exc:  # noqa: BLE001
            raise RuntimeError(f"XLSX dependency missing: {exc}") from exc

        wb = Workbook()
        default_ws = wb.active

        sheets = document.get("sheets")
        if isinstance(sheets, dict) and sheets:
            wb.remove(default_ws)
            for name, rows in sheets.items():
                ws = wb.create_sheet(str(name))
                if not isinstance(rows, list):
                    raise ValueError("each value in document.sheets must be rows list")
                for row in rows:
                    if not isinstance(row, list):
                        raise ValueError("each row in document.sheets must be a list")
                    ws.append([str(cell) for cell in row])
        else:
            rows = document.get("rows")
            if not isinstance(rows, list):
                raise ValueError("document.rows must be a list for xlsx")
            sheet_name = str(document.get("sheet") or "Sheet1")
            default_ws.title = sheet_name
            for row in rows:
                if not isinstance(row, list):
                    raise ValueError("each row in document.rows must be a list")
                default_ws.append([str(cell) for cell in row])
        wb.save(str(path))

    def _generate_pdf(self, path: Path, *, blocks: list[Any]) -> None:
        try:
            from reportlab.lib.pagesizes import A4  # pyright: ignore[reportMissingImports,reportMissingModuleSource]
            from reportlab.pdfgen import canvas  # pyright: ignore[reportMissingImports,reportMissingModuleSource]
        except Exception as exc:  # noqa: BLE001
            raise RuntimeError(f"PDF dependency missing: {exc}") from exc

        c = canvas.Canvas(str(path), pagesize=A4)
        width, height = A4
        x = 50
        y = height - 50
        for block in blocks:
            if y < 70:
                c.showPage()
                y = height - 50
            if not isinstance(block, dict):
                continue
            block_type = str(block.get("type") or "paragraph")
            text = str(block.get("text") or "")
            if not text:
                continue
            if block_type == "heading":
                c.setFont("Helvetica-Bold", 14)
                c.drawString(x, y, text[:1000])
                y -= 24
            elif block_type == "table":
                c.setFont("Helvetica", 10)
                rows = block.get("rows")
                if isinstance(rows, list):
                    for row in rows:
                        row_text = " | ".join(str(cell) for cell in (row if isinstance(row, list) else [row]))
                        c.drawString(x, y, row_text[:1000])
                        y -= 16
                else:
                    c.drawString(x, y, text[:1000])
                    y -= 16
                y -= 4
            else:
                c.setFont("Helvetica", 11)
                for line in text.splitlines() or [text]:
                    c.drawString(x, y, line[:1000])
                    y -= 16
                y -= 4
        c.save()


TOOL = DocumentGenerateTool
