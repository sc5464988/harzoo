"""Structured document editing tool for machine consumption."""


from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Any

import yaml

from harzoo.agent.kernel.tool import Tool, ToolResult

TOOL_VERSION = "2026-05-22"


def _resolve_path(path: str) -> Path:
    p = Path(path).expanduser()
    if p.is_absolute():
        return p.resolve()
    return (Path.cwd() / p).resolve()


class DocumentEditTool(Tool):
    """结构化文档编辑工具，按操作指令修改文档而非自由文本替换。"""

    name = "DocumentEdit"
    danger_level = 1
    description = "Edit structured documents with explicit operations for common formats."
    parameters = {
        "type": "object",
        "properties": {
            "file_path": {"type": "string", "description": "Target file path"},
            "format": {"type": "string", "enum": ["txt", "md", "json", "yaml", "xml", "html", "csv", "xlsx", "docx"]},
            "operation": {"type": "string", "enum": ["overwrite", "append", "upsert_rows"]},
            "data": {"description": "Operation payload"},
            "sheet": {"type": "string", "description": "Sheet name for xlsx operations"},
            "create_if_missing": {"type": "boolean", "default": True},
            "encoding": {"type": "string", "default": "utf-8"},
        },
        "required": ["file_path", "format", "operation", "data"],
    }
    risk_level = "medium"
    side_effect = True
    idempotent = False

    def execute(
        self,
        file_path: str,
        format: str,
        operation: str,
        data: Any,
        sheet: str | None = None,
        create_if_missing: bool = True,
        encoding: str = "utf-8",
        **_: Any,
    ) -> ToolResult:
        fmt = str(format).strip().lower()
        op = str(operation).strip().lower()
        if fmt not in {"txt", "md", "json", "yaml", "xml", "html", "csv", "xlsx", "docx"}:
            return ToolResult.failure(f"Unsupported format: {fmt}", code="UNSUPPORTED_FORMAT")
        if op not in {"overwrite", "append", "upsert_rows"}:
            return ToolResult.failure(f"Unsupported operation: {op}", code="UNSUPPORTED_OPERATION")

        path = _resolve_path(file_path)
        if not path.exists() and not create_if_missing:
            return ToolResult.failure(f"File not found: {file_path}", code="FILE_NOT_FOUND")

        try:
            path.parent.mkdir(parents=True, exist_ok=True)
            if fmt in {"txt", "md", "xml", "html"}:
                return self._edit_text(path, op, data, fmt=fmt, encoding=encoding)
            if fmt == "json":
                return self._edit_json(path, op, data, encoding=encoding)
            if fmt == "yaml":
                return self._edit_yaml(path, op, data, encoding=encoding)
            if fmt == "csv":
                return self._edit_csv(path, op, data, encoding=encoding)
            if fmt == "xlsx":
                return self._edit_xlsx(path, op, data, sheet=sheet)
            if fmt == "docx":
                return self._edit_docx(path, op, data)
            return ToolResult.failure(f"Unsupported format: {fmt}", code="UNSUPPORTED_FORMAT")
        except Exception as exc:  # noqa: BLE001
            return ToolResult.failure(f"{type(exc).__name__}: {exc}", code="WRITE_FAILED")

    def _edit_text(self, path: Path, op: str, data: Any, *, fmt: str, encoding: str) -> ToolResult:
        content = str(data)
        if op == "overwrite":
            path.write_text(content, encoding=encoding)
        elif op == "append":
            previous = path.read_text(encoding=encoding, errors="replace") if path.exists() else ""
            path.write_text(previous + content, encoding=encoding)
        else:
            return ToolResult.failure(f"Operation {op} not supported for {fmt}", code="UNSUPPORTED_OPERATION")
        return ToolResult.success(
            {
                "path": str(path),
                "format": fmt,
                "operation": op,
                "bytes_written": path.stat().st_size,
                "warnings": [],
            }
        )

    def _edit_json(self, path: Path, op: str, data: Any, *, encoding: str) -> ToolResult:
        if op == "append":
            return ToolResult.failure("append is not supported for json", code="UNSUPPORTED_OPERATION")
        if op == "overwrite":
            payload = data
        elif op == "upsert_rows":
            existing = {}
            if path.exists():
                existing_raw = json.loads(path.read_text(encoding=encoding, errors="replace"))
                if not isinstance(existing_raw, dict):
                    return ToolResult.failure("upsert_rows for json requires existing object", code="INVALID_ARGUMENTS")
                existing = existing_raw
            if not isinstance(data, dict):
                return ToolResult.failure("upsert_rows for json requires object data", code="INVALID_ARGUMENTS")
            merged = dict(existing)
            merged.update(data)
            payload = merged
        else:
            return ToolResult.failure(f"Operation {op} not supported for json", code="UNSUPPORTED_OPERATION")
        path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding=encoding)
        return ToolResult.success(
            {
                "path": str(path),
                "format": "json",
                "operation": op,
                "bytes_written": path.stat().st_size,
                "warnings": [],
            }
        )

    def _edit_yaml(self, path: Path, op: str, data: Any, *, encoding: str) -> ToolResult:
        if op == "append":
            return ToolResult.failure("append is not supported for yaml", code="UNSUPPORTED_OPERATION")
        if op == "overwrite":
            payload = data
        elif op == "upsert_rows":
            existing = {}
            if path.exists():
                existing_raw = yaml.safe_load(path.read_text(encoding=encoding, errors="replace"))
                if existing_raw is None:
                    existing_raw = {}
                if not isinstance(existing_raw, dict):
                    return ToolResult.failure("upsert_rows for yaml requires existing object", code="INVALID_ARGUMENTS")
                existing = existing_raw
            if not isinstance(data, dict):
                return ToolResult.failure("upsert_rows for yaml requires object data", code="INVALID_ARGUMENTS")
            merged = dict(existing)
            merged.update(data)
            payload = merged
        else:
            return ToolResult.failure(f"Operation {op} not supported for yaml", code="UNSUPPORTED_OPERATION")
        path.write_text(yaml.safe_dump(payload, allow_unicode=True, sort_keys=False), encoding=encoding)
        return ToolResult.success(
            {
                "path": str(path),
                "format": "yaml",
                "operation": op,
                "bytes_written": path.stat().st_size,
                "warnings": [],
            }
        )

    def _normalize_rows(self, data: Any) -> list[list[str]]:
        if not isinstance(data, list):
            raise ValueError("rows data must be a list")
        rows: list[list[str]] = []
        for row in data:
            if not isinstance(row, list):
                raise ValueError("each row must be a list")
            rows.append([str(cell) for cell in row])
        return rows

    def _edit_csv(self, path: Path, op: str, data: Any, *, encoding: str) -> ToolResult:
        if op == "upsert_rows":
            return ToolResult.failure(
                "upsert_rows is only supported for json and yaml; use overwrite or append for csv",
                code="UNSUPPORTED_OPERATION",
            )
        rows = self._normalize_rows(data)
        if op not in {"overwrite", "append", "upsert_rows"}:
            return ToolResult.failure(f"Operation {op} not supported for csv", code="UNSUPPORTED_OPERATION")
        mode = "w" if op == "overwrite" else "a"
        with path.open(mode, encoding=encoding, newline="") as fh:
            writer = csv.writer(fh)
            writer.writerows(rows)
        return ToolResult.success(
            {
                "path": str(path),
                "format": "csv",
                "operation": op,
                "bytes_written": path.stat().st_size,
                "rows_written": len(rows),
                "warnings": [],
            }
        )

    def _edit_xlsx(self, path: Path, op: str, data: Any, *, sheet: str | None) -> ToolResult:
        if op == "upsert_rows":
            return ToolResult.failure(
                "upsert_rows is only supported for json and yaml; use overwrite or append for xlsx",
                code="UNSUPPORTED_OPERATION",
            )
        try:
            from openpyxl import Workbook, load_workbook  # pyright: ignore[reportMissingImports,reportMissingModuleSource]
        except Exception as exc:  # noqa: BLE001
            return ToolResult.failure(f"XLSX dependency missing: {exc}", code="PARSE_FAILED")
        rows = self._normalize_rows(data)
        target_sheet = str(sheet or "Sheet1")

        if path.exists():
            wb = load_workbook(str(path))
        else:
            wb = Workbook()
        if target_sheet in wb.sheetnames:
            ws = wb[target_sheet]
            if op == "overwrite":
                ws.delete_rows(1, ws.max_row)
        else:
            ws = wb.create_sheet(target_sheet)

        if op not in {"overwrite", "append", "upsert_rows"}:
            return ToolResult.failure(f"Operation {op} not supported for xlsx", code="UNSUPPORTED_OPERATION")
        for row in rows:
            ws.append(row)
        wb.save(str(path))
        return ToolResult.success(
            {
                "path": str(path),
                "format": "xlsx",
                "operation": op,
                "sheet": target_sheet,
                "bytes_written": path.stat().st_size,
                "rows_written": len(rows),
                "warnings": [],
            }
        )

    def _edit_docx(self, path: Path, op: str, data: Any) -> ToolResult:
        if op not in {"overwrite", "append"}:
            return ToolResult.failure(f"Operation {op} not supported for docx", code="UNSUPPORTED_OPERATION")
        try:
            from docx import Document  # pyright: ignore[reportMissingImports]
        except Exception as exc:  # noqa: BLE001
            return ToolResult.failure(f"DOCX dependency missing: {exc}", code="PARSE_FAILED")

        existed = path.exists()
        if op == "overwrite":
            doc = Document()
        elif existed:
            doc = Document(str(path))
        else:
            doc = Document()
        text = str(data)
        for line in text.splitlines() or [text]:
            if line.strip():
                doc.add_paragraph(line)
        doc.save(str(path))
        warnings: list[str] = []
        if op == "overwrite" and existed:
            warnings.append("overwrite replaced the entire document")
        return ToolResult.success(
            {
                "path": str(path),
                "format": "docx",
                "operation": op,
                "bytes_written": path.stat().st_size,
                "warnings": warnings,
            }
        )


TOOL = DocumentEditTool
